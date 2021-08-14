#pdf translator
from PyPDF2 import pdf
from tqdm import tqdm,trange
from time import sleep
from colorama import Fore,init
from requests import post
from random import random
from pdf2docx import Converter
from docx import Document
from os.path import isfile
from os import remove
from docx2pdf import convert
from argparse import ArgumentParser
from random import random
import threading
init()
class translate:
    def __init__(self,document,lenguaje,to_traduce):
        self.lenguaje = lenguaje
        self.to_traduce = to_traduce
    def block_partition(self,it,block_size):
        if len(it)>5000:
            return [it[x:x+block_size] for x in range(0,len(it),block_size)]
        else:
            return [it]
    def req(self,it:"text"):
        for text in self.block_partition(it,120):
            data = {"sl":self.lenguaje,"tl":self.to_traduce,"q":text}
            headers = {"Charset":"UTF-8","User-Agent":"AndroidTranslate/5.3.0.RC02.130475354-53000263 5.1 phone TRANSLATE_OPM5_TEST_1"}
            url = "https://translate.google.com/translate_a/single?client=at&dt=t&dt=ld&dt=qca&dt=rm&dt=bd&dj=1&hl=es-ES&ie=UTF-8&oe=UTF-8&inputm=2&otf=2&iid=1dd3b944-fa62-4b55-b330-74909a99969e"
            resp = post(url,data=data,headers=headers)
            tme = random()
            #print( Fore.CYAN + f"tiempo de espera:{tme}".center(50,"="))
            sleep(tme)
            try:
                assert resp.raise_for_status() != 200
                res = resp.json()
                for x in res.get('sentences',''):
                    if x.get("trans",None) == None:
                        print(f"{Fore.RED} no se ha podido traducir el texto:\n {text} \n {Fore.YELLOW} a;adiendo texto sin traducir.")
                    yield x.get("trans",text)
            except AssertionError:
                print(f"{Fore.RED} No se ha podido traducir el texto {text}")
class read(translate):
    def __init__(self,document,documentOut,lenguaje,to_traduce,UseMP):
        super().__init__(document,lenguaje,to_traduce)
        self._documentNameOut = documentOut
        self.document = document
        self.UseMP = UseMP
        self.flgMP = self.validate()
        self.convert2docx()
    def validate(self):
        if self.UseMP == None:
            self.UseMP = 1
            return False
        else:
            return True
    @property
    def documentNameOut(self):
        return self._documentNameOut
    def compute(self)->"docx":
        self.pdf = open(self.document,"rb")
        self.open = pdf.PdfFileReader(self.pdf)
        print(Fore.CYAN + "informacion del documento".center(50,"="))
        try:
            for x in self.open.getDocumentInfo().items():
                print(f"{x[0]}: {x[1]}")
        except:
            print(f"{Fore.RED}No se ha podido obtener la informacion")
        print(Fore.BLUE + "fin".center(50,"="))
        try:
            #return self.docxExtractText()
            for x,n in self.docxExtractText():
                for y in super().req(x):
                    self.docxLoad.paragraphs[n].text = y
        except Exception as e:
            print(f"{Fore.RED} Ha ocurrido un error: {e}")
        self.docxLoad.save(self.docx)
        return self
    def __enter__(self):
        '''
        if self.UseTreads:
            thread = threading.Thread(target=self.compute)
            thread.start()
        else:
            self.compute()'''
        self.compute()

    def convert2docx(self):
        print(Fore.CYAN)
        print("convirtiendo".center(50,'-'))
        self.docx = self.delExtension(self.document) + 'out.docx'
        self.c = Converter(self.document)
        self.c.convert(self.docx,multi_processing=self.flgMP,cpu_count=self.UseMP)
        print(Fore.CYAN)
        print("fin".center(50,'-'))
        print('\n\n\n\n\n')
        sleep(4)
    def docxExtractText(self):
        #pdf -> docx
        self.docxLoad = Document(self.docx)
        with tqdm(total=len(self.docxLoad.paragraphs),colour='green') as pbar:
            pbar.set_description(desc="traduciendo")
            for x,number in zip(self.docxLoad.paragraphs,range(0,len(self.docxLoad.paragraphs) - 1)):
                #sleep
                pbar.update(1)
                yield x.text,number
            #self.subTH = threading.Thread(target=self.subTH,args=(pbar))
    def nextText(self)->str:
        #!unused TH
        # pdf -> text
        try:
            with tqdm(total=self.open.getNumPages(),colour="green") as pbar:
                pbar.set_description(desc="traduciendo")
                for i in range(0,self.open.getNumPages()):
                    #sleep(0.2)
                    pg = self.open.getPage(i)
                    pbar.update(1)
                    yield pg.extractText(),i
        except Exception as e:
            print(f"{Fore.RED}Ha ocurrido un error en la lectura del documento".center(50,"/"))
            print(f"{Fore.RED} {e}")
    @property
    def docOut(self):
        return self._documentNameOut
    @docOut.setter
    def docOut(self,arg):
        self._documentNameOut = arg
    def save2pdf(self):
        pass
    def __exit__(self,tipo_excepcion,valor_excepcion,traza_error):
        print(Fore.GREEN)
        self.pdf.close()
    def convert2pdf(self):
        print("conversion a pdf".center(50,'-'))
        try:
            convert(self.docx,self._documentNameOut)
            remove(self.docx)
        except Exception as e:
            print(f"{Fore.RED}Ha ocurrido una excepcion {e}")
        finally:
            print("fin".center(50,'-'))
    @staticmethod
    def delExtension(ext):
        return "".join([ x for x in ext.split('.') if x != 'pdf'])

def edDocx(par):
    #
    with read(par.path,par.output,par.languaje,par.traduce,par.UseMultiProcess) as tr:
        if par.NotConvert2pdf:
            pass
        if par.Convert2pdf:
            tr.convert2pdf()
        print(f"{Fore.CYAN} tu archivo lo puedes encontrar en: {tr.documentNameOut}")
        print(f"{Fore.GREEN}" + "fin".center(50,'-'))
def main():
    ar = ArgumentParser()
    ar.add_argument('path',help="el archivo pdf")
    ar.add_argument('-o','--output',help="especificar el nombre de un archivo de salida")
    ar.add_argument('-noC','--NotConvert2pdf',help=" deolver el archivo como un .docx opcion por defecto",action='store_true')
    ar.add_argument('-c','--Convert2pdf',help="convertir el archivo a pdf despues de traducirlo como .docx",action='store_true')
    ar.add_argument('-l','--languaje',help="el lenguje en el que esta el documento")
    ar.add_argument('-t','--traduce',help='lenguaje a traducir')
    ar.add_argument('-um','--UseMultiProcess',help="usar multiprocesos para una traduccion mas rapida colocar los cpu a usar",type=int)
    par = ar.parse_args()
    if not(par.languaje) or not(par.traduce):
        print(f"{Fore.RED} No se ha podido traducir \n {Fore.BLUE}[{Fore.YELLOW}*{Fore.BLUE}]      faltan argumentos")
        return 0
    if par.output:
        edDocx(par)
    else:
        par.output = read.delExtension(par.path) + f'{random()}__.pdf'
        print(f"{Fore.CYAN} tu archivo lo puedes encontrar como: {par.output}")
        edDocx(par)
if __name__ == '__main__':
    main()